"""Read a file from the user's document folders.

Ported from zylch/services/solve_tools.py::_read_document. Honors the
platform-aware path logic (commit 87e6216) and does NOT truncate extracted
text — the full content is returned to the LLM.

Native extraction for the common office formats (PDF, DOCX, XLSX) is done
in-process so the LLM doesn't have to fall back to ``run_python`` — that
tool is in APPROVAL_TOOLS and would force a "Conferma richiesta" prompt
on the user for a harmless read. The dependencies (``pypdf``,
``python-docx``, ``openpyxl``) are declared in ``pyproject.toml``; if any
of them is missing at runtime the corresponding format degrades to the
old "Use run_python to process this file" hint instead of crashing.
"""

import glob
import logging
import os
from typing import Any, Dict, List

from .base import Tool, ToolResult, ToolStatus

logger = logging.getLogger(__name__)


# Plain-text extensions we can read directly.
TEXT_EXTS = (".txt", ".md", ".csv", ".json", ".xml", ".log", ".yaml", ".yml")


def _extract_pdf_text(path: str) -> str:
    """Return the full text of a PDF, page by page, via pypdf.

    Raises ImportError if pypdf is not installed; the caller falls back
    to the legacy "use run_python" hint in that case.
    """
    from pypdf import PdfReader  # local import — keeps tool importable without pypdf

    reader = PdfReader(path)
    parts: List[str] = []
    for idx, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as e:
            logger.warning(f"[read_document] pypdf page {idx} extract failed: {e}")
            page_text = ""
        parts.append(f"--- page {idx} ---\n{page_text}")
    return "\n\n".join(parts)


def _extract_docx_text(path: str) -> str:
    """Return the text of a .docx file (paragraphs + table cells)."""
    from docx import Document  # local import — keeps tool importable without python-docx

    doc = Document(path)
    out: List[str] = []
    for para in doc.paragraphs:
        if para.text:
            out.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                out.append("\t".join(cells))
    return "\n".join(out)


def _extract_xlsx_text(path: str) -> str:
    """Return the contents of an .xlsx workbook, sheet by sheet, as TSV-ish blocks."""
    from openpyxl import load_workbook  # local import — keeps tool importable without openpyxl

    wb = load_workbook(path, data_only=True, read_only=True)
    sections: List[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: List[str] = []
        for row in ws.iter_rows(values_only=True):
            if not row:
                continue
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                rows.append("\t".join(cells))
        sections.append(f"--- sheet: {sheet_name} ---\n" + "\n".join(rows))
    return "\n\n".join(sections)


def _collect_search_paths() -> List[str]:
    """Platform-aware document search paths.

    `~/Downloads` is ALWAYS included (implicitly), even when DOCUMENT_PATHS
    is set — that's where `download_attachment` places its output, so
    `read_document` must be able to find those files without extra config.
    """
    home = os.path.expanduser("~")
    profile_dir = os.environ.get("ZYLCH_PROFILE_DIR", "")
    # Downloads folder: respect the user's DOWNLOADS_DIR override from
    # Settings — download_attachment uses the same variable, so we stay
    # in sync without another knob.
    configured_downloads = os.environ.get("DOWNLOADS_DIR", "").strip()
    downloads = (
        os.path.expanduser(configured_downloads)
        if configured_downloads
        else os.path.join(home, "Downloads")
    )

    defaults = [
        os.path.join(home, "gdrive-shared"),
        os.path.join(home, "Documents"),
        downloads,
        "/tmp/zylch/attachments",
        "/tmp/zylch",
    ]
    if profile_dir:
        defaults.append(profile_dir)

    doc_paths = os.environ.get("DOCUMENT_PATHS", "")
    if doc_paths:
        configured = [os.path.expanduser(p.strip()) for p in doc_paths.split(",") if p.strip()]
        paths = [p for p in configured if os.path.isdir(p)]
        if not paths:
            paths = [p for p in defaults if os.path.isdir(p)]
    else:
        paths = [p for p in defaults if os.path.isdir(p)]

    # Ensure ~/Downloads is always in the search set (idempotent).
    if os.path.isdir(downloads) and downloads not in paths:
        paths.append(downloads)

    return paths


class ReadDocumentTool(Tool):
    """Read a file from the user's document folders."""

    def __init__(self):
        super().__init__(
            name="read_document",
            description=(
                "Read a file from the user's document folders."
                " Searches by filename across all registered paths."
                " Accepts absolute paths too."
            ),
        )

    async def execute(self, filename: str = "", **kwargs) -> ToolResult:
        logger.debug(f"[read_document] execute(args={{'filename_len':" f" {len(filename)}}})")
        if not filename:
            result = ToolResult(
                status=ToolStatus.ERROR,
                data=None,
                error="No filename provided",
            )
            logger.debug(f"[read_document] -> status={result.status}")
            return result

        # Absolute path shortcut
        path: str = ""
        if os.path.isabs(filename) and os.path.isfile(filename):
            path = filename
        else:
            paths = _collect_search_paths()
            if not paths:
                result = ToolResult(
                    status=ToolStatus.ERROR,
                    data=None,
                    error=("No document folders found." " Add DOCUMENT_PATHS to your profile .env"),
                )
                logger.debug(f"[read_document] -> status={result.status}")
                return result

            found: List[str] = []
            for base in paths:
                pattern = os.path.join(base, "**", f"*{filename}*")
                found.extend(glob.glob(pattern, recursive=True))

            # If filename has a path component, try it as a relative path too.
            if not found and os.path.isfile(filename):
                found.append(os.path.abspath(filename))

            if not found:
                result = ToolResult(
                    status=ToolStatus.ERROR,
                    data=None,
                    error=(f"No file matching '{filename}' in:" f" {', '.join(paths)}"),
                )
                logger.debug(f"[read_document] -> status={result.status}")
                return result

            path = found[0]

        ext = os.path.splitext(path)[1].lower()

        if ext == ".pdf":
            try:
                text = _extract_pdf_text(path)
                result = ToolResult(
                    status=ToolStatus.SUCCESS,
                    data={"text": text, "path": path, "format": "pdf"},
                    message=f"PDF: {path}\n\n{text}",
                )
                logger.debug(
                    f"[read_document] -> status={result.status} format=pdf bytes={len(text)}"
                )
                return result
            except ImportError:
                result = ToolResult(
                    status=ToolStatus.SUCCESS,
                    data={"text": "", "path": path, "format": "pdf"},
                    message=(
                        f"Found PDF: {path}\n"
                        f"pypdf not installed — fall back to run_python to read it."
                    ),
                )
                logger.warning(
                    f"[read_document] pypdf missing, returning stub for {path}"
                )
                return result
            except Exception as e:
                logger.error(f"[read_document] pdf extract failed: {e}")
                result = ToolResult(
                    status=ToolStatus.ERROR,
                    data=None,
                    error=f"Could not read {path}: {e}",
                )
                return result

        if ext == ".docx":
            try:
                text = _extract_docx_text(path)
                result = ToolResult(
                    status=ToolStatus.SUCCESS,
                    data={"text": text, "path": path, "format": "docx"},
                    message=f"DOCX: {path}\n\n{text}",
                )
                logger.debug(
                    f"[read_document] -> status={result.status} format=docx bytes={len(text)}"
                )
                return result
            except ImportError:
                result = ToolResult(
                    status=ToolStatus.SUCCESS,
                    data={"text": "", "path": path, "format": "docx"},
                    message=(
                        f"Found DOCX: {path}\n"
                        f"python-docx not installed — fall back to run_python."
                    ),
                )
                logger.warning(f"[read_document] python-docx missing, returning stub for {path}")
                return result
            except Exception as e:
                logger.error(f"[read_document] docx extract failed: {e}")
                result = ToolResult(
                    status=ToolStatus.ERROR,
                    data=None,
                    error=f"Could not read {path}: {e}",
                )
                return result

        if ext in (".xlsx", ".xlsm"):
            try:
                text = _extract_xlsx_text(path)
                result = ToolResult(
                    status=ToolStatus.SUCCESS,
                    data={"text": text, "path": path, "format": ext.lstrip(".")},
                    message=f"XLSX: {path}\n\n{text}",
                )
                logger.debug(
                    f"[read_document] -> status={result.status} format=xlsx bytes={len(text)}"
                )
                return result
            except ImportError:
                result = ToolResult(
                    status=ToolStatus.SUCCESS,
                    data={"text": "", "path": path, "format": ext.lstrip(".")},
                    message=(
                        f"Found {ext.upper()}: {path}\n"
                        f"openpyxl not installed — fall back to run_python."
                    ),
                )
                logger.warning(f"[read_document] openpyxl missing, returning stub for {path}")
                return result
            except Exception as e:
                logger.error(f"[read_document] xlsx extract failed: {e}")
                result = ToolResult(
                    status=ToolStatus.ERROR,
                    data=None,
                    error=f"Could not read {path}: {e}",
                )
                return result

        if ext in TEXT_EXTS:
            try:
                with open(path, "r", errors="replace") as f:
                    content = f.read()
                result = ToolResult(
                    status=ToolStatus.SUCCESS,
                    data={
                        "text": content,
                        "path": path,
                        "format": ext.lstrip(".") or "text",
                    },
                    message=f"File: {path}",
                )
                logger.debug(f"[read_document] -> status={result.status}" f" bytes={len(content)}")
                return result
            except Exception as e:
                logger.error(f"[read_document] read failed: {e}")
                result = ToolResult(
                    status=ToolStatus.ERROR,
                    data=None,
                    error=f"Could not read {path}: {e}",
                )
                logger.debug(f"[read_document] -> status={result.status}")
                return result

        result = ToolResult(
            status=ToolStatus.SUCCESS,
            data={
                "text": "",
                "path": path,
                "format": ext.lstrip(".") or "binary",
            },
            message=(f"Found: {path} ({ext})\n" f"Use run_python to process this file."),
        )
        logger.debug(f"[read_document] -> status={result.status} format={ext}")
        return result

    def get_schema(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "description": self.description,
            "input_schema": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": ("Filename or partial name to search"),
                    },
                },
                "required": ["filename"],
            },
        }
