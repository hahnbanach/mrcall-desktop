"""Interactive chat mode for Zylch standalone.

REPL with readline (history, line editing) and rich (markdown output).
Routes slash commands to command_handlers, natural language to ChatService.
"""

import asyncio
import logging
import os
import readline
import shlex
import sys

from rich.console import Console
from rich.markdown import Markdown

from zylch.cli.utils import get_owner_id, load_env
from zylch.tools.config import ToolConfig

logger = logging.getLogger(__name__)

console = Console()

# Commands that need ToolConfig with BYOK credentials
_BYOK_COMMANDS = {
    "/memory", "/email", "/agent", "/calendar",
}
# Commands that need only (args, owner_id)
_SIMPLE_COMMANDS = {
    "/tasks", "/stats", "/jobs", "/reset", "/tutorial",
}


def _setup_readline(profile_dir: str):
    """Configure readline: history file, tab completion, bindings."""
    history_file = os.path.join(profile_dir, ".zylch_history")

    # Load existing history (may fail on macOS sandbox)
    try:
        readline.read_history_file(history_file)
    except (FileNotFoundError, PermissionError, OSError):
        pass

    try:
        readline.set_history_length(1000)
        import atexit
        atexit.register(readline.write_history_file, history_file)
    except (PermissionError, OSError):
        pass  # Read-only or sandboxed — no history

    # Slash command completion
    commands = [
        "/help", "/quit", "/update", "/sync",
        "/tasks", "/stats",
        "/email", "/memory", "/agent", "/calendar",
        "/connect", "/mrcall", "/jobs", "/clear",
        "/model",
    ]

    def completer(text, state):
        if text.startswith("/"):
            matches = [c for c in commands if c.startswith(text)]
        else:
            matches = []
        return matches[state] if state < len(matches) else None

    readline.set_completer(completer)
    readline.parse_and_bind("tab: complete")


def _print_dashboard(profile: str, owner_id: str):
    """Show startup dashboard: profile, sync state, pending work."""
    from datetime import datetime, timezone
    from zylch.storage.storage import Storage

    store = Storage.get_instance()

    # --- Email stats ---
    email_stats = store.get_email_stats(owner_id)
    total_emails = email_stats.get("total_emails", 0)

    # Last update = from SyncState, fallback to newest email date
    last_sync_str = "never"
    sync_age_hours = None

    sync_state = store.get_sync_state(owner_id)
    last_update = None
    if sync_state:
        last_update = sync_state.get("last_dream_at")
    if not last_update:
        last_update = email_stats.get("latest_date")

    if last_update:
        try:
            if isinstance(last_update, str):
                ls = datetime.fromisoformat(last_update)
            else:
                ls = last_update
            if ls.tzinfo is None:
                ls = ls.replace(tzinfo=timezone.utc)
            delta = datetime.now(timezone.utc) - ls
            sync_age_hours = delta.total_seconds() / 3600
            if delta.days > 0:
                last_sync_str = f"{delta.days}d ago"
            elif delta.seconds >= 3600:
                last_sync_str = (
                    f"{delta.seconds // 3600}h ago"
                )
            elif delta.seconds >= 60:
                last_sync_str = (
                    f"{delta.seconds // 60}m ago"
                )
            else:
                last_sync_str = "just now"
        except Exception:
            last_sync_str = "unknown"

    # --- Pending processing ---
    try:
        pending_memory = len(
            store.get_unprocessed_emails(owner_id)
        )
    except Exception:
        pending_memory = 0
    try:
        pending_tasks = len(
            store.get_unprocessed_emails_for_task(owner_id)
        )
    except Exception:
        pending_tasks = 0

    # --- Active tasks ---
    task_stats = store.get_task_items_stats(owner_id)
    if task_stats:
        active = (
            task_stats["action_required"]
            - task_stats["completed"]
        )
        if active < 0:
            active = 0
    else:
        active = 0

    # --- Task urgency breakdown ---
    urgency_line = ""
    if active > 0:
        try:
            items = store.get_task_items(
                owner_id, action_required=True, limit=200,
            )
            by_urgency = {}
            for t in items:
                u = (t.get("urgency") or "unknown").lower()
                by_urgency[u] = by_urgency.get(u, 0) + 1
            parts = []
            for u in ["critical", "high", "medium", "low"]:
                if by_urgency.get(u):
                    parts.append(f"{by_urgency[u]} {u}")
            if parts:
                urgency_line = f" ({', '.join(parts)})"
        except Exception:
            pass

    # --- Memory ---
    try:
        from zylch.memory.blob_storage import BlobStorage

        mem_stats = BlobStorage().get_stats(owner_id)
        entities = mem_stats.get("total_blobs", 0)
    except Exception:
        entities = 0

    # --- Print ---
    from zylch import __version__

    # Time-aware greeting
    from datetime import datetime
    hour = datetime.now().hour
    if hour < 12:
        greeting = "Good morning"
    elif hour < 18:
        greeting = "Good afternoon"
    else:
        greeting = "Good evening"

    # User name from env or profile
    user_name = os.environ.get("USER_FULL_NAME", "")
    if user_name:
        user_name = user_name.split()[0]  # First name
    else:
        user_name = profile.split("@")[0]  # mario.alemi → mario

    console.print()
    console.print(
        f"[bold cyan]{greeting}, {user_name}![/bold cyan]"
        f" [dim](Zylch v{__version__})[/dim]",
    )
    console.print()

    # Context-aware intro
    if total_emails == 0:
        console.print(
            "  I'm Zylch, your sales intelligence agent.\n"
            "  I connect to your email, WhatsApp, and phone\n"
            "  to find what needs your attention — and help\n"
            "  you get it done.\n\n"
            "  Let's start by syncing your messages.\n",
        )
    elif active > 0:
        console.print(
            f"  I've been keeping an eye on your messages.\n"
            f"  You have {active} things that need attention.\n"
            f"  I can help you handle them — review, reply,\n"
            f"  or solve them step by step.\n",
        )
    else:
        console.print(
            "  All caught up! I can search your contacts,\n"
            "  draft emails, or look for new messages.\n",
        )

    # Status line
    if total_emails > 0:
        console.print(
            f"  {total_emails} emails synced"
            f" (last: {last_sync_str})",
        )
    else:
        console.print(
            "  No emails synced yet.",
        )

    if entities > 0:
        console.print(
            f"  {entities} contacts in memory",
        )

    if active > 0:
        console.print(
            f"  {active} tasks needing action{urgency_line}",
        )

    # --- Menu ---
    console.print()

    console.print(
        "  [bold]What would you like to do?[/bold]",
    )
    console.print()

    options = []
    idx = 1
    if active > 0:
        options.append(
            (str(idx), f"Show your {active} tasks", "/tasks interactive"),
        )
        idx += 1
    if total_emails == 0 or (
        sync_age_hours is not None and sync_age_hours > 0.5
    ) or pending_memory or pending_tasks:
        age = last_sync_str if total_emails > 0 else "never"
        options.append(
            (str(idx), f"Look for new messages (last update: {age})", "/update"),
        )
        idx += 1
    options.append(
        (str(idx), "Let's chat!", None),
    )

    for key, label, _ in options:
        console.print(f"  [bold]{key})[/bold] {label}")
    console.print()

    # Get choice
    while True:
        try:
            choice = input("  > ").strip().lower()
        except (EOFError, KeyboardInterrupt):
            choice = "c"
            break
        if choice in [o[0] for o in options]:
            break
        if not choice:
            choice = options[0][0]  # Default to first
            break

    # Execute choice
    for key, _, cmd in options:
        if choice == key and cmd:
            return cmd  # Return command to execute

    return None  # Chat mode

    console.print()


def _print_response(text: str):
    """Render response as rich markdown."""
    if not text:
        return
    try:
        console.print(Markdown(text))
    except Exception:
        # Fallback to plain text if rich fails
        print(text)


def interactive_chat():
    """REPL-style chat loop with Zylch AI.

    - Selects profile (if multiple exist)
    - Acquires exclusive lock on profile
    - readline for history + line editing
    - rich for markdown output
    - Lines starting with / are dispatched as slash commands.
    - Everything else goes to ChatService.process_message().
    - Ctrl-C or /quit exits.
    """
    from zylch.cli.profiles import (
        get_active_profile,
        get_active_profile_dir,
    )

    # Profile already selected and activated by main.py _setup_profile()
    load_env()
    owner_id = get_owner_id()

    # Setup readline with profile-specific history
    profile_dir = get_active_profile_dir()
    if profile_dir:
        _setup_readline(profile_dir)

    # Reap zombie jobs from dead sessions
    from zylch.storage.storage import Storage
    zombies = Storage.get_instance().reap_zombie_jobs()
    if zombies:
        console.print(
            f"[yellow]Cleaned {zombies} failed job(s)"
            f" from previous session.[/yellow]\n"
        )

    profile = get_active_profile()
    logger.info(
        f"[chat] Starting interactive chat,"
        f" profile={profile}, owner_id={owner_id}"
    )

    startup_cmd = _print_dashboard(profile, owner_id)
    console.print(
        "[dim]Ctrl+C interrupts, Ctrl+D twice to exit.[/dim]\n"
    )

    conversation_history: list = []
    _last_eof = False  # For double Ctrl+D exit

    # Execute startup choice if any
    if startup_cmd:
        _handle_slash_command(
            startup_cmd, owner_id, conversation_history,
        )
        console.print()

    while True:
        try:
            user_input = input("you> ").strip()
            _last_eof = False
        except EOFError:
            if _last_eof:
                console.print("\n[dim]Bye![/dim]")
                sys.exit(0)
            _last_eof = True
            console.print(
                "\n[dim]Press Ctrl+D again to exit.[/dim]"
            )
            continue
        except KeyboardInterrupt:
            _last_eof = False
            console.print()  # Clean line after ^C
            continue

        if not user_input:
            continue

        if user_input.lower() in ("/quit", "/exit", "/q"):
            console.print("[dim]Bye![/dim]")
            sys.exit(0)

        try:
            if user_input.startswith("/"):
                _handle_slash_command(
                    user_input, owner_id,
                    conversation_history,
                )
            else:
                # Attach file content if path detected
                user_input = _expand_file_refs(user_input)
                _handle_chat_message(
                    user_input, owner_id,
                    conversation_history,
                )
        except KeyboardInterrupt:
            console.print(
                "\n[yellow]Interrupted.[/yellow]"
            )


def _expand_file_refs(text: str) -> str:
    """Detect file paths in user input and attach content.

    Supports: /path/to/file, ~/path, paths with backslash escapes.
    For .docx/.pdf, uses python-docx/pypdf to extract text.
    """
    import re

    # Match paths: /... or ~/... (with optional backslash escapes)
    path_pattern = r'(?:/[\w./@-]+(?:\\ [\w./@-]+)*|~[\w./@-]+(?:\\ [\w./@-]+)*)'
    matches = re.findall(path_pattern, text)
    if not matches:
        return text

    for match in matches:
        clean = match.replace("\\ ", " ").strip("'\"")
        expanded = os.path.expanduser(clean)

        if not os.path.isfile(expanded):
            continue

        ext = os.path.splitext(expanded)[1].lower()
        content = None

        try:
            if ext in (".txt", ".md", ".csv", ".json", ".xml"):
                with open(expanded, "r", errors="replace") as f:
                    content = f.read()

            elif ext == ".docx":
                try:
                    import docx
                    doc = docx.Document(expanded)
                    content = "\n".join(
                        p.text for p in doc.paragraphs
                    )
                except ImportError:
                    import subprocess
                    result = subprocess.run(
                        ["python", "-c",
                         f"import docx; doc=docx.Document('{expanded}');"
                         f"print('\\n'.join(p.text for p in doc.paragraphs))"],
                        capture_output=True, text=True, timeout=10,
                    )
                    if result.returncode == 0:
                        content = result.stdout

            elif ext == ".pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(expanded)
                    content = "\n".join(
                        page.extract_text() or ""
                        for page in reader.pages
                    )
                except ImportError:
                    pass

            else:
                # Binary file — just note it exists
                size = os.path.getsize(expanded)
                text += (
                    f"\n\n[File attached: {os.path.basename(expanded)}"
                    f" ({size} bytes, {ext})]"
                )
                continue

        except Exception as e:
            text += (
                f"\n\n[Could not read {os.path.basename(expanded)}: {e}]"
            )
            continue

        if content:
            fname = os.path.basename(expanded)
            console.print(
                f"  [dim]Attached: {fname}"
                f" ({len(content)} chars)[/dim]",
            )
            text += (
                f"\n\n--- FILE: {fname} ---\n"
                f"{content}\n--- END FILE ---"
            )

    return text


def _handle_agent_run(
    instructions: str,
    owner_id: str,
    conversation_history: list,
):
    """Run agentic loop with all tools (from chat)."""
    from zylch.api.token_storage import get_active_llm_provider
    from zylch.llm.client import LLMClient
    from zylch.services.task_interactive import (
        APPROVAL_TOOLS,
        SOLVE_SYSTEM_PROMPT,
        SOLVE_TOOLS,
        _format_approval_preview,
        _get_personal_data_section,
        _run_agent_loop,
    )
    from zylch.services.solve_tools import execute_tool
    from zylch.storage.storage import Storage

    provider, api_key = get_active_llm_provider(owner_id)
    if not api_key:
        console.print("[red]No API key. Run zylch init.[/red]")
        return

    user_email = os.environ.get("EMAIL_ADDRESS", "")
    user_name = (
        user_email.split("@")[0] if user_email else "you"
    )
    store = Storage.get_instance()

    client = LLMClient(api_key=api_key, provider=provider)
    system = SOLVE_SYSTEM_PROMPT.format(
        user_name=user_name,
        personal_data_section=_get_personal_data_section(),
    )

    if not instructions:
        console.print(
            "  What should I do?"
            " (type, empty line to send)",
        )
        lines = []
        while True:
            try:
                line = input("  > ")
            except (EOFError, KeyboardInterrupt):
                return
            if not line:
                break
            lines.append(line)
        instructions = "\n".join(lines)
        instructions = _expand_file_refs(instructions)
        if not instructions.strip():
            return

    messages = [
        {"role": "user", "content": instructions},
    ]

    console.print("\n  [dim]Working...[/dim]")
    messages = _run_agent_loop(
        client, system, messages, store, owner_id,
    )

    # Post-loop: continue conversation or done
    while True:
        console.print(
            "\n  [bold]d)[/bold] Done"
            "   or type to continue"
            " (empty line to send)",
        )
        lines = []
        while True:
            try:
                line = input("  > ")
            except (EOFError, KeyboardInterrupt):
                return
            if not line and lines:
                break
            if not line and not lines:
                return  # Enter = done
            lines.append(line)

        followup = "\n".join(lines).strip()
        followup = _expand_file_refs(followup)
        if followup.lower() == "d":
            return

        messages.append(
            {"role": "user", "content": followup},
        )
        console.print("\n  [dim]Working...[/dim]")
        messages = _run_agent_loop(
            client, system, messages, store, owner_id,
        )


def _handle_slash_command(
    raw_input: str,
    owner_id: str,
    conversation_history: list,
):
    """Dispatch a slash command to the appropriate handler."""
    try:
        parts = shlex.split(raw_input)
    except ValueError as e:
        console.print(f"[red]Error: malformed command — {e}[/red]")
        return

    cmd = parts[0].lower()
    args = parts[1:] if len(parts) > 1 else []
    logger.debug(f"[chat] slash cmd={cmd}, args={args}")

    from zylch.services.command_handlers import (
        COMMAND_HANDLERS,
    )

    if cmd == "/clear":
        conversation_history.clear()
        console.print("[dim]Conversation cleared.[/dim]")
        return

    # /agent run <instructions> — agentic loop with tools
    if cmd == "/agent" and args and args[0] == "run":
        instructions = " ".join(args[1:]) if len(args) > 1 else ""
        instructions = _expand_file_refs(instructions)
        _handle_agent_run(
            instructions, owner_id, conversation_history,
        )
        return

    if cmd not in COMMAND_HANDLERS:
        console.print(
            f"[yellow]Unknown command: {cmd}. Type /help.[/yellow]"
        )
        return

    handler = COMMAND_HANDLERS[cmd]

    try:
        if cmd in ("/sync", "/update"):
            config = ToolConfig.from_settings()
            result = asyncio.run(
                handler(args, config, owner_id)
            )
        elif cmd in _BYOK_COMMANDS:
            config = ToolConfig.from_settings_with_owner(
                owner_id
            )
            if cmd == "/agent":
                ctx = {
                    "_conversation_history": (
                        conversation_history
                    ),
                }
                result = asyncio.run(
                    handler(args, config, owner_id, ctx)
                )
            else:
                result = asyncio.run(
                    handler(args, config, owner_id)
                )
        elif cmd in _SIMPLE_COMMANDS:
            result = asyncio.run(handler(args, owner_id))
        elif cmd in ("/help", "/echo"):
            if args:
                result = asyncio.run(handler(args))
            else:
                result = asyncio.run(handler())
        elif cmd in ("/share", "/revoke", "/connect"):
            result = asyncio.run(
                handler(args, owner_id, None)
            )
        elif cmd == "/mrcall":
            result = asyncio.run(
                handler(args, owner_id, None, None)
            )
        else:
            result = asyncio.run(handler(args, owner_id))

        if result:
            _print_response(result)

    except Exception as e:
        logger.error(
            f"[chat] Command {cmd} failed: {e}",
            exc_info=True,
        )
        console.print(f"[red]Error running {cmd}: {e}[/red]")


def _handle_chat_message(
    user_input: str,
    owner_id: str,
    conversation_history: list,
):
    """Send a natural-language message to ChatService."""
    from zylch.services.chat_service import ChatService

    logger.debug(
        f"[chat] Sending to ChatService: "
        f"{repr(user_input[:80])}"
    )

    try:
        service = _get_chat_service()
        result = asyncio.run(
            service.process_message(
                user_message=user_input,
                user_id=owner_id,
                conversation_history=conversation_history,
                context={"user_id": owner_id},
            )
        )

        response = result.get("response", "")
        if response:
            console.print()
            _print_response(response)
            console.print()

        # Update history for next turn
        conversation_history.append(
            {"role": "user", "content": user_input}
        )
        conversation_history.append(
            {"role": "assistant", "content": response}
        )

    except Exception as e:
        logger.error(
            f"[chat] ChatService error: {e}", exc_info=True
        )
        console.print(f"[red]Error: {e}[/red]")


# Lazy singleton for ChatService
_chat_service_instance = None


def _get_chat_service():
    """Get or create the ChatService singleton."""
    global _chat_service_instance
    if _chat_service_instance is None:
        from zylch.services.chat_service import ChatService

        _chat_service_instance = ChatService()
    return _chat_service_instance
